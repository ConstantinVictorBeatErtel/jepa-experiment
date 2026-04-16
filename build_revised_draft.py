from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/ConstiX/Downloads/JEPA")
OUTPUT = Path("/Users/ConstiX/Downloads/INDENG242B_ Mini_Project_vf_revised.docx")
FIGURE = ROOT / "results" / "viz" / "summary_best4.png"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "World Models - Revisiting Feature Prediction for Learning Visual Representations from Video "
        "(Bardes et al., 2024), grounded in I-JEPA (Assran et al., 2023)"
    )
    run.bold = True
    run.font.size = Pt(15)

    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qrun = quote.add_run(
        '"Pure logical thinking cannot yield us any knowledge of the empirical world; '
        'all knowledge of reality starts from experience and ends in it." '
        "— Albert Einstein, Ideas and Opinions (1954)"
    )
    qrun.italic = True

    doc.add_heading("I. Introduction", level=1)
    doc.add_paragraph(
        "Abstract (why this paper was chosen and why it matters): V-JEPA "
        "(Revisiting Feature Prediction for Learning Visual Representations from Video, Bardes et al., 2024) "
        "is a self-supervised video representation learning method from Meta FAIR. Its significance extends beyond "
        "benchmark performance: it is one of the clearest empirical tests of Yann LeCun’s claim that intelligence "
        "should be trained to predict abstract world state rather than reconstruct raw observations. The central JEPA "
        "idea is to predict latent representations of missing content from visible context, thereby pushing the model "
        "toward structure, objects, and relationships instead of surface-level texture. Because the experiment in this "
        "project is image-based rather than video-based, the directly relevant methodological reference is the earlier "
        "I-JEPA paper (Assran et al., 2023), which V-JEPA extends from images to spatiotemporal data. This revised "
        "draft therefore treats V-JEPA as the motivating world-model paper while tying the actual replication to the "
        "original image JEPA design that we implemented."
    )

    doc.add_heading("II. Summary of Main Technical Content and Contributions", level=1)
    doc.add_paragraph(
        "From Attention to JEPA: The technical lineage begins with the transformer of Vaswani et al. (2017), which "
        "showed that self-attention can model long-range dependencies without recurrence. Dosovitskiy et al. (2020) "
        "adapted this idea to images with Vision Transformers by treating image patches as tokens. MAE (He et al., 2022) "
        "then showed that masking and reconstructing pixels can produce scalable self-supervised pretraining. However, "
        "pixel reconstruction also forces the encoder to model fine local detail that is not always aligned with semantic "
        "understanding. I-JEPA (Assran et al., 2023) proposes a different target space: instead of reconstructing pixels, "
        "it predicts the latent representation of missing blocks from visible context. V-JEPA (Bardes et al., 2024) "
        "extends the same idea from images to videos by replacing image patches with spatiotemporal tubelets and by "
        "learning predictive structure across time."
    )
    doc.add_paragraph(
        "Technical explanation: At the image level, I-JEPA tokenizes an image into non-overlapping patches and uses a "
        "Vision Transformer as the context encoder. The context encoder only processes visible context patches; masked "
        "tokens are excluded entirely, rather than replaced by a [MASK] token. A second encoder, the target encoder, "
        "processes the full image and produces patch-level target representations. Its parameters are not updated by "
        "backpropagation but by exponential moving average (EMA) of the context encoder. A narrow predictor transformer "
        "takes the context representation plus learnable mask-position tokens and predicts the latent representation of "
        "each target block. The loss is the average squared L2 distance between predicted and target representations over "
        "the masked target patches."
    )
    doc.add_paragraph(
        "A core insight of the I-JEPA paper is that the masking strategy is not cosmetic. The paper’s proposed "
        "multi-block masking strategy samples four relatively large target blocks and one informative context block. "
        "In Table 6 of Assran et al. (2023), this multi-block masking strategy dramatically outperforms rasterized, "
        "single-block, and random masking on ImageNet-1% linear evaluation. The paper also shows, in Table 7, that "
        "predicting target-encoder representations is much better than predicting pixels, which directly supports the "
        "latent-prediction hypothesis. In Appendix C, the paper further shows that sufficiently large target blocks, a "
        "large context block, several targets, and a narrow predictor are all important for good semantics."
    )
    doc.add_paragraph(
        "Main contributions: I-JEPA demonstrates that semantic visual features can be learned without hand-crafted view "
        "augmentations by predicting in representation space. V-JEPA then extends this principle to video and world-modeling. "
        "Together, the two papers make a coherent argument: if the objective is to learn structure and prediction rather than "
        "surface statistics, latent prediction is often a better inductive bias than raw pixel reconstruction."
    )

    doc.add_heading("III. Relevance to the Course", level=1)
    doc.add_paragraph(
        "JEPA connects directly to several themes from INDENG 242B. First, it is a clean example of how the learning "
        "objective shapes what backpropagation discovers: changing the target space from pixels to latent features changes "
        "the semantic level of the learned representation. Second, it highlights the transformer as a general-purpose "
        "architecture for structured prediction beyond language. Third, the EMA target encoder is closely related to the "
        "course discussion of optimization dynamics and stability: rather than updating both branches symmetrically, JEPA "
        "stabilizes the target network with a slow-moving teacher, which helps avoid collapse. Finally, JEPA clarifies an "
        "important modeling tradeoff discussed in class: a model can be highly accurate at reconstructing local detail yet "
        "still learn a weaker representation for downstream semantic tasks."
    )

    doc.add_heading("IV. Implementation/Replication", level=1)
    doc.add_paragraph(
        "This replication is best understood as an image-based, resource-constrained adaptation of I-JEPA rather than a "
        "full reproduction of V-JEPA. The experiment uses CIFAR-10 static images, not video, so it cannot test temporal "
        "reasoning, motion, or causality. What it can test is the narrower design claim shared by I-JEPA and V-JEPA: "
        "whether predicting latent target representations yields more semantic features than reconstructing masked pixels "
        "when the backbone and masking budget are otherwise controlled."
    )
    doc.add_paragraph(
        "Implementation details: CIFAR-10 images were resized on the fly to 96×96 and normalized with standard CIFAR-10 "
        "statistics. Both models used separate small ViT encoders with patch size 8, embedding dimension 256, depth 6, "
        "4 attention heads, MLP ratio 4.0, dropout 0.1, sinusoidal 2D positional embeddings, and no [CLS] token. In direct "
        "accordance with I-JEPA Appendix A.1, evaluation used average pooling over patch tokens, and JEPA evaluation used "
        "the target encoder rather than the context encoder."
    )
    doc.add_paragraph(
        "Masking was made deliberately close to the original I-JEPA image setup. For each image, four possibly overlapping "
        "target blocks were sampled with scale in the range (0.15, 0.2) of total patches and aspect ratio in the range "
        "(0.75, 1.5). One context block was sampled with scale in the range (0.85, 1.0) and unit aspect ratio, and any "
        "overlap with target patches was removed. This mirrors the masking described in Figure 4 and Appendix A.1 of "
        "Assran et al. (2023). As in the paper, the context and target masks were produced by a collator in the data "
        "loader, and patch counts were equalized within a batch for efficient processing."
    )
    doc.add_paragraph(
        "The JEPA model used a context encoder, an EMA target encoder, and a narrow 6-layer predictor with a 128-dimensional "
        "bottleneck. The target encoder was initialized identically to the context encoder, updated with EMA using momentum "
        "0.996 linearly increased to 1.0, and never optimized by gradient descent. The MAE baseline used the same encoder "
        "family but reconstructed normalized pixel patches on the union of the masked targets with a lightweight 2-layer "
        "transformer decoder. Both models were trained for 100 epochs using AdamW, linear warmup from 1e-4 to 1e-3, cosine "
        "decay to 1e-6, weight decay increased from 0.04 to 0.4, and gradient clipping at norm 1.0."
    )
    doc.add_paragraph(
        "This revised implementation supersedes the earlier draft. In the earlier draft, MAE appeared stronger than JEPA. "
        "That conclusion is no longer supported by the final experiment. The direction changed because the final pipeline "
        "implemented the key I-JEPA ingredients much more faithfully: the exact multi-block masking geometry, full-image "
        "EMA teacher targets, evaluation with the target encoder, exclusion of masked tokens from the encoder input, and "
        "a narrow predictor bottleneck. This reversal is consistent with the original paper’s own ablations. Assran et al. "
        "show that representation-space targets, large semantic target blocks, informative context, and predictor design "
        "substantially affect downstream performance; MAE is comparatively robust to imperfect setup because pixel "
        "reconstruction is an easier local objective."
    )

    doc.add_paragraph(
        "Table 1 reports the main quantitative results in a format intentionally similar to the transfer tables in the "
        "I-JEPA paper."
    )
    add_caption(
        doc,
        "Table 1. Linear-probe transfer and retrieval on CIFAR-10. Following the style of Assran et al. (2023), "
        "we report frozen-representation performance for methods without view data augmentations."
    )
    table1 = doc.add_table(rows=4, cols=5)
    table1.style = "Table Grid"
    headers = ["Method", "Arch.", "Linear Probe Top-1", "k-NN acc@1", "k-NN acc@5"]
    for idx, header in enumerate(headers):
        set_cell_text(table1.rows[0].cells[idx], header, bold=True)
    set_repeat_table_header(table1.rows[0])
    set_cell_text(table1.rows[1].cells[0], "Methods without view data augmentations", bold=True)
    for idx in range(1, 5):
        set_cell_text(table1.rows[1].cells[idx], "")
    row = table1.rows[2].cells
    set_cell_text(row[0], "MAE")
    set_cell_text(row[1], "ViT-256d / p8")
    set_cell_text(row[2], "56.29%")
    set_cell_text(row[3], "50.59%")
    set_cell_text(row[4], "83.59%")
    row = table1.rows[3].cells
    set_cell_text(row[0], "I-JEPA")
    set_cell_text(row[1], "ViT-256d / p8")
    set_cell_text(row[2], "71.08%")
    set_cell_text(row[3], "66.15%")
    set_cell_text(row[4], "89.13%")

    doc.add_paragraph(
        "The final outcome is clear: JEPA outperformed MAE by 14.79 percentage points on linear probe, by 15.56 points on "
        "k-NN acc@1, and by 5.54 points on k-NN acc@5. This direction now matches the original I-JEPA paper, which reports "
        "that latent target prediction outperforms pixel reconstruction on semantic transfer tasks. In other words, once the "
        "experiment was aligned with the original design, the static-image CIFAR-10 result moved toward the same qualitative "
        "conclusion as Assran et al. (2023)."
    )

    doc.add_paragraph(
        "Table 2 makes the relation to the original paper explicit by listing the most important design choices that were "
        "copied from I-JEPA."
    )
    add_caption(
        doc,
        "Table 2. Alignment between this replication and the original I-JEPA design of Assran et al. (2023)."
    )
    table2 = doc.add_table(rows=9, cols=3)
    table2.style = "Table Grid"
    headers = ["Design choice", "I-JEPA paper", "This experiment"]
    for idx, header in enumerate(headers):
        set_cell_text(table2.rows[0].cells[idx], header, bold=True)
    set_repeat_table_header(table2.rows[0])
    rows = [
        ("No [CLS] token", "Yes", "Yes"),
        ("Average-pool patch tokens for eval", "Yes", "Yes"),
        ("Use target encoder for JEPA evaluation", "Yes", "Yes"),
        ("4 target blocks", "Yes", "Yes"),
        ("Target scale (0.15, 0.2)", "Yes", "Yes"),
        ("Context scale (0.85, 1.0)", "Yes", "Yes"),
        ("EMA target encoder 0.996 → 1.0", "Yes", "Yes"),
        ("Narrow predictor bottleneck", "Yes", "Yes (128 vs 256 encoder dim)"),
    ]
    for ridx, values in enumerate(rows, start=1):
        for cidx, value in enumerate(values):
            set_cell_text(table2.rows[ridx].cells[cidx], value)

    doc.add_paragraph(
        "Qualitatively, the nearest-neighbor retrieval analysis also aligned with the paper’s semantic-prediction story. "
        "The average JEPA semantic retrieval score over 20 selected test queries was 0.502, and the qualitative examples "
        "showed that JEPA predictor outputs retrieved training images with stronger class and object-level consistency than "
        "the corresponding MAE features. This is similar in spirit to the paper’s qualitative analysis of what information "
        "is preserved by the predictor: JEPA keeps object-level structure while discarding some exact low-level detail."
    )

    if FIGURE.exists():
        doc.add_picture(str(FIGURE), width=Inches(6.4))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(
            doc,
            "Figure 1. Qualitative nearest-neighbor retrieval summary from the final CIFAR-10 experiment. As in the "
            "I-JEPA paper’s qualitative analysis, the key pattern is not exact pixel match but semantic consistency: "
            "JEPA predictor retrievals are more object- and class-aligned than MAE retrievals."
        )

    doc.add_heading("V. Possible Extensions and Novel Applications", level=1)
    doc.add_paragraph(
        "One promising extension of V-JEPA would be to make the predictor action-conditioned for robotics and Physical AI. "
        "In that setting, the latent variable conditioning the predictor would represent robot actions or control inputs, "
        "allowing the model to predict future latent scene states several steps ahead. This would move the system from "
        "passive representation learning toward planning and counterfactual prediction."
    )
    doc.add_paragraph(
        "A second extension would apply JEPA-style latent prediction to anomaly detection in industrial or medical settings. "
        "Instead of measuring anomaly score by pixel reconstruction error, one could measure latent prediction error relative "
        "to normal scene structure. Because the revised experiment shows that JEPA is more semantically organized than MAE, "
        "this extension is better motivated as a semantic-structure detector than in the earlier draft. However, the current "
        "project did not directly evaluate anomaly detection, so this should be framed as a hypothesis for future work rather "
        "than as an empirical result of the present experiment."
    )

    doc.add_heading("Works Cited", level=1)
    works = [
        'Assran, Mahmoud, et al. "Self-Supervised Learning from Images with a Joint-Embedding Predictive Architecture." '
        "arXiv preprint arXiv:2301.08243, 2023.",
        'Bardes, Adrien, et al. "Revisiting Feature Prediction for Learning Visual Representations from Video." '
        "arXiv preprint arXiv:2404.08471, 2024.",
        'Vaswani, Ashish, et al. "Attention Is All You Need." Advances in Neural Information Processing Systems, 2017.',
        'Dosovitskiy, Alexey, et al. "An Image Is Worth 16x16 Words: Transformers for Image Recognition at Scale." '
        "arXiv preprint arXiv:2010.11929, 2020.",
        'He, Kaiming, et al. "Masked Autoencoders Are Scalable Vision Learners." Proceedings of the IEEE/CVF Conference '
        "on Computer Vision and Pattern Recognition (CVPR), 2022.",
        'LeCun, Yann. "A Path Towards Autonomous Machine Intelligence." OpenReview, 2022.',
        'Bardes, Adrien, et al. "V-JEPA 2: Self-supervised video models enable understanding, prediction and planning." '
        "arXiv preprint arXiv:2506.09985, 2025.",
    ]
    for item in works:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_heading("Appendix", level=1)
    doc.add_paragraph("Github: https://github.com/ConstantinVictorBeatErtel/jepa-experiment")
    doc.add_paragraph(
        "Experiment Results: The earlier draft’s appendix was left blank and did not reflect the final run. The table below "
        "records the actual outputs of the completed 100-epoch experiment."
    )
    add_caption(
        doc,
        "Table A1. Final experiment results from the completed CIFAR-10 run."
    )
    table3 = doc.add_table(rows=3, cols=4)
    table3.style = "Table Grid"
    headers = ["Method", "Linear Probe Top-1", "k-NN acc@1", "k-NN acc@5"]
    for idx, header in enumerate(headers):
        set_cell_text(table3.rows[0].cells[idx], header, bold=True)
    set_repeat_table_header(table3.rows[0])
    row = table3.rows[1].cells
    set_cell_text(row[0], "I-JEPA")
    set_cell_text(row[1], "71.08%")
    set_cell_text(row[2], "66.15%")
    set_cell_text(row[3], "89.13%")
    row = table3.rows[2].cells
    set_cell_text(row[0], "MAE")
    set_cell_text(row[1], "56.29%")
    set_cell_text(row[2], "50.59%")
    set_cell_text(row[3], "83.59%")

    doc.add_paragraph(
        "Overall interpretation: The final experiment supports the same qualitative lesson emphasized by the original I-JEPA "
        "paper. When the masking strategy, target encoder, and predictor are implemented in a way that encourages semantic "
        "prediction rather than local reconstruction shortcuts, JEPA-style latent prediction can produce substantially stronger "
        "frozen representations than an MAE-style pixel decoder, even in a small static-image setting."
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
